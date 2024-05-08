import os


import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def createInvitations(txtFile, docName):
    """Creates invitations based on names in txt file
    Args:
        txtFile (str): text file to read from
        docName (str): doc file to save invitations in
    """
    doc = docx.Document()

    intro = "It would be a pleasure to have the company of"
    address = "at 11101 Memory lane on the evening of"
    date = "April 31st"
    time = "at 24 O'Clock"

    with open(txtFile) as guestList:
        for guest in guestList:
            name = guest[:-1]
            p1 = doc.add_paragraph()
            p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            f1 = p1.add_run(intro)
            f1.font.bold = True
            f1.font.italic = True
            f1.font.size = Pt(13)

            p2 = doc.add_paragraph()
            p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            f2 = p2.add_run(name)
            f2.font.bold = True
            f2.font.size = Pt(15)

            p3 = doc.add_paragraph()
            p3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            f3 = p3.add_run(address)
            f3.font.bold = True
            f3.font.italic = True
            f3.font.size = Pt(12)

            p4 = doc.add_paragraph()
            p4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            f4 = p4.add_run(date)
            f4.font.size = Pt(12)

            p5 = doc.add_paragraph()
            p5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            f5 = p5.add_run(time)
            f5.font.bold = True
            f5.font.italic = True
            f5.font.size = Pt(12)

            doc.add_page_break()

    doc.save(docName)


if __name__ == "__main__":
    createInvitations("guests.txt", "invitations.docx")
